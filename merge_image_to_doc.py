# -*- encoding: utf-8 -*-
"""
-------------------------------------------------
   File Name：    merge_image_to_doc.py
   Description :
   Author :       Wings DH
   Time：         2020/11/8 2:07 下午
-------------------------------------------------
   Change Activity:
                   2020/11/8: Create
-------------------------------------------------
"""
import os
import sys
from docx import Document
import re

from docx.shared import Inches


def get_sec_pos(image_name):
    min_pos, sec_pos = re.findall('\d+', image_name)
    return int(min_pos) * 60 + int(sec_pos)


def main(argv=None):
    if argv is None:
        argv = sys.argv

    src_d_path = './output'
    tgt_doc_f_path = './merge.docx'

    images = [name for name in os.listdir(src_d_path) if name.endswith('jpg')]
    images.sort(key=get_sec_pos)
    images = [os.path.join(src_d_path, name) for name in images]

    doc = Document()
    for image in images:
        doc.add_paragraph(image)
        doc.add_picture(image, width=Inches(5.0))
    doc.save(tgt_doc_f_path)


if __name__ == "__main__":
    sys.exit(main())
